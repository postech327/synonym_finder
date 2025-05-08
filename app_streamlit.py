import streamlit as st
from openai import OpenAI  # ✅ 최신 SDK
import os
import json
import pandas as pd
from dotenv import load_dotenv
from prompts import make_prompt
from utils import parse_table
from docx import Document  # Word 저장용
from datetime import datetime  # 파일명 자동화용
from gtts import gTTS  # ✅ 음성 저장용 (gTTS는 영어/한국어 지원)

# .env에서 API 키 불러오기
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))  # ✅ 최신 방식

SAVE_FILE = "saved_words.json"
RESULT_FILE = "saved_results.json"

# 🔹 저장된 단어 목록 불러오기
def load_saved_words():
    if os.path.exists(SAVE_FILE):
        with open(SAVE_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

# 🔹 유의어 결과 불러오기
def load_saved_results():
    if os.path.exists(RESULT_FILE):
        with open(RESULT_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

# 🔹 단어 저장
def save_user_word(username, word):
    data = load_saved_words()
    if username not in data:
        data[username] = []
    if word not in data[username]:
        data[username].append(word)
        with open(SAVE_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

# 🔹 유의어 결과 저장
def save_user_results(username, result):
    data = load_saved_results()
    if username not in data:
        data[username] = []
    data[username].extend(result)
    with open(RESULT_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# 🔹 단어 불러오기
def get_user_words(username):
    data = load_saved_words()
    return data.get(username, [])

# 🔹 결과 불러오기
def get_user_results(username):
    data = load_saved_results()
    return data.get(username, [])

# 🔹 분석용 프롬프트 생성
def make_analysis_prompt(text):
    return f"""
This GPT has two main functions: flow analysis and question generation, both based on a given passage. Users provide a text passage, and the GPT will analyze its flow and generate relevant multiple-choice questions in English and Korean. All summaries provided will be concise and in Korean, limited to less than 20 characters.

**Flow Analysis (in Korean)**
1. The passage is broken down into three sections—Introduction, Body, and Conclusion.
2. For each section, the GPT will:
  - **Subject Matter**: Identify and describe the main focus of the section in Korean.
  - **Summary**: Provide a brief summary in Korean, ensuring it is less than 20 characters.

**Question Generation (Both in English and Korean)**
1. **Multiple-Choice Questions**: Create relevant questions with 5 answer choices. Each question will have:
  - **Title**: A question to determine the best title for the passage, using styles and examples based on those in the attached reference PDF. Ensure options are of varying sentence lengths and meanings for clarity and distinction.
  - **Topic**: A question to determine the best topic for the passage, following similar guidelines as the title question. On Topic, choices should be given in small letters.
  - **Main Idea**: A question to determine the best main idea for the passage, with 5 choices listed by sentence length. The correct answer will be distinct, and the wrong answers will differ in content and context to avoid overlapping meaning with the correct one.
  - **Summary**: Summarize the passage in one sentence.

Each question will be provided with one correct answer, and the answers will be translated into Korean. The wrong answers should be translated into Korean as well. Make the choices in circled digits like ① ② ③ ④ ⑤. English and Korean choices should be given separately. To ensure variety, the correct answer should not always be option ①, but instead randomly assigned among the five options for each question.
📌 You **must provide the correct answer explicitly** for each question in both English and Korean.
Passage:
{text}
    """

# 🔹 불일치 문제용 프롬프트 생성
def make_false_statements_prompt(text):
    return f"""
This GPT generates multiple-choice questions (MCQs) in both English and Korean based on a given passage. It follows a structured format where:

1. A passage is provided by the user.
2. A question is created asking the user to select all choices that do not match the passage.
3. Answer choices are labeled with ⓐ ⓑ ⓒ ⓓ ⓔ ⓕ ⓖ ⓗ.
4. The wording of answer choices is varied to avoid direct repetition of the passage.
5. Just three incorrect choices are included.(They are the answers: only three)
6. Correct answers are provided with detailed explanations.
7. The entire question, choices, answers, and explanations are provided in both **English and Korean**.
8. English and Korean should be separated.

This GPT is ideal for language learners, educators, and test creators who need bilingual comprehension questions. It ensures accuracy and consistency while maintaining clarity in both languages.

    Passage:
    {text}
"""

# 🔹 빈칸 문제용 프롬프트 생성
def make_blank_question_prompt(text):
    return f"""
This GPT generates multiple-choice questions (MCQs) in both English and Korean based on a given passage. It follows a structured format where: 

1. A passage is provided by the user.
2. A question is created by turning the topic sentence or key idea of the passage into a fill-in-the-blank format.
3. Answer choices are labeled with ① ② ③ ④ ⑤. 
4. The correct answer must match the original sentence from the passage exactly, while the incorrect choices are created by modifying the correct answer to express opposite or irrelevant ideas.
5. The entire question, choices, and answers are provided in both **English and Korean**. 
6. English and Korean should be separated.

Passage:
{text}
"""

# 🔹 Word 저장 함수
def save_to_word_file(content, filename=None):
    doc = Document()
    doc.add_heading("GPT 문제 생성 결과", level=1)
    for line in content.split("\n"):
        doc.add_paragraph(line)
    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"gpt_questions_{timestamp}.docx"
    doc.save(filename)
    return filename

# 🔹 MP3 저장 함수 (gTTS 이용)
st.sidebar.markdown("---")
st.sidebar.subheader("🎧 텍스트 음성 변환")
custom_voice_input = st.sidebar.text_area("📜 음성으로 변환할 텍스트를 입력하세요")
voice_lang = st.sidebar.radio("🌐 음성 언어 선택", ["한국어", "영어"], horizontal=True)

if st.sidebar.button("🔊 음성 저장 (mp3)"):
    if not custom_voice_input.strip():
        st.sidebar.warning("텍스트를 입력해주세요.")
    else:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"tts_output_{timestamp}.mp3"

        lang_code = "ko" if voice_lang == "한국어" else "en"
        tts = gTTS(text=custom_voice_input, lang=lang_code)
        tts.save(filename)

        with open(filename, "rb") as f:
            st.sidebar.download_button(
                label="📥 다운로드: 음성 파일",
                data=f,
                file_name=filename,
                mime="audio/mpeg"
            )

# 🔹 Streamlit UI 인터페이스
st.title("📘 GPT 기반 영어 문제 생성기")

menu = st.sidebar.radio("기능 선택", ["주제·제목·요지", "불일치 문제", "빈칸 문제", "전체 문제 생성"])
input_text = st.text_area("✏️ 지문을 입력하세요", height=300)

if st.button("🎯 문제 생성"):
    if not input_text.strip():
        st.warning("지문을 입력해주세요.")
    else:
        prompts = {
            "주제·제목·요지": make_analysis_prompt(input_text),
            "불일치 문제": make_false_statements_prompt(input_text),
            "빈칸 문제": make_blank_question_prompt(input_text),
        }
        if menu == "전체 문제 생성":
            combined_result = ""
            with st.spinner("GPT로부터 응답을 기다리는 중..."):
                for title, prompt in prompts.items():
                    response = client.chat.completions.create(
                        model="gpt-4",
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.7
                    )
                    content = response.choices[0].message.content
                    combined_result += f"🔷 {title} 결과\n\n{content}\n\n{'='*80}\n\n"
            st.session_state.generated_content = combined_result
            st.success("✅ 전체 문제 생성 완료!")
            st.text_area("🧾 전체 결과", value=combined_result, height=500)
        else:
            prompt = prompts[menu]
            with st.spinner("GPT로부터 응답을 기다리는 중..."):
                response = client.chat.completions.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.7
                )
                content = response.choices[0].message.content
                st.session_state.generated_content = content
                st.success("✅ 생성 완료!")
                st.text_area("🧾 생성된 결과", value=content, height=400)

# 🔹 저장 버튼
if 'generated_content' in st.session_state and st.session_state.generated_content:
    col1, col2 = st.columns(2)
    with col1:
        if st.button("📄 Word로 저장"):
            filename = save_to_word_file(st.session_state.generated_content)
            with open(filename, "rb") as f:
                st.download_button(
                    label="📥 다운로드: Word 파일",
                    data=f,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    with col2:
        if st.button("🔊 음성(mp3)으로 저장"):
            filename = save_to_mp3(st.session_state.generated_content)
            with open(filename, "rb") as f:
                st.download_button(
                    label="📥 다운로드: 음성 파일",
                    data=f,
                    file_name=filename,
                    mime="audio/mpeg"
                )
